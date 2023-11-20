
import chardet
from docx import Document
import shutil
import os
from copy import deepcopy

from docx.shared import Inches


def copyDoc(number_act):

    source_file = 'C:/Users/Нур/PycharmProjects/CreaterDocx/creater/docxer/temp/template.docx'
    source_Da = 'C:/Users/Нур/PycharmProjects/CreaterDocx/creater/docxer/temp/docxDA.docx'

    destination_file_Temp = f'C:/Users/Нур/PycharmProjects/CreaterDocx/creater/docxer/acts/{number_act}'
    destination_file_Da = f'C:/Users/Нур/PycharmProjects/CreaterDocx/creater/docxer/acts/{number_act}'


    path = f'C:/Users/Нур/PycharmProjects/CreaterDocx/creater/docxer/acts/{number_act}/template.docx'
    pathDa = f'C:/Users/Нур/PycharmProjects/CreaterDocx/creater/docxer/acts/{number_act}/docxDA.docx'


    if(os.path.exists(path) == False):
        shutil.copy2(source_file, destination_file_Temp)
    else:
        print("Файл уже существует ")
    if(os.path.exists(pathDa) == False):
        shutil.copy2(source_Da,destination_file_Da)


def openDocx(number_act):
    path = f'C:/Users/Нур/PycharmProjects/CreaterDocx/creater/docxer/acts/{number_act}/docxDA.docx'
    if (os.path.exists(path)):
        os.startfile(path)
    else:print("Файл не существует ")



def append_date(numact, kpp, inn, nameZaivitel, adress, OGRN,adreesPlaceProizvod):
    # Открой шаблонный docx-файл
    template_path = f'C:/Users/Нур/PycharmProjects/CreaterDocx/creater/docxer/acts/{numact}/template.docx'
    tagDa = "{{da}}"
    template = Docugit initment(template_path)
    tableDaPath = f'C:/Users/Нур/PycharmProjects/CreaterDocx/creater/docxer/acts/{numact}/docxDA.docx'
    tableDa = Document(tableDaPath)
    table = tableDa.tables[0]

    # Замени теги на данные пользователя
    for index, paragraph in enumerate(template.paragraphs):
        for run in paragraph.runs:
            # Получи текст run
            text = run.text
            # Замени теги на данные пользователя
            text = text.replace('{{NUM}}',numact)
            text = text.replace('{{name_zaivitel}}', nameZaivitel)
            text = text.replace('{{inn}}', inn)
            text = text.replace('{{kpp}}', kpp)
            text = text.replace('{{OGRN}}', OGRN)
            text = text.replace('{{adress}}', adress)
            text = text.replace('{{adress_place}}',adreesPlaceProizvod)

            # Установи новый текст run
            run.text = text
        if '{da}' in paragraph.text:
            # Remove the existing paragraph with {da}
            run = paragraph.runs[0]
            paragraph._p.remove(run._r)
            # Insert a new table
            tableinTemplDa = template.add_table(rows=len(table.rows), cols=len(table.columns))
            print('aadd')
            rows_list = list(table.rows)
            print(rows_list)

            for i, row in enumerate(rows_list):
                for j, cell in enumerate(row.cells):
                    tableinTemplDa.cell(i, j).text = cell.text
            tableinTemplDa.autofit = False



            # Insert the table into the document
            paragraph._p.addnext(tableinTemplDa._tbl)

        # Save the new docx file



    template.save(template_path)

    print("All funcution works right")

    return True


def append_files(numact, TableDA):
    awaitFilesPath = f'C:/Users/Нур/PycharmProjects/CreaterDocx/creater/docxer/acts/{numact}/docxDA.docx'
    doc = Document(TableDA)
    doc.save(awaitFilesPath)
    return True

def createAct(numAct):
    directory = 'C:/Users/Нур/PycharmProjects/CreaterDocx/creater/docxer/acts'
    print(numAct)

    # Получаем список существующих папок
    # Создаем новую папку
    new_folder_name = str(numAct)
    new_folder_path = os.path.join(directory, new_folder_name)
    os.mkdir(new_folder_path)
    copyDoc(numAct)